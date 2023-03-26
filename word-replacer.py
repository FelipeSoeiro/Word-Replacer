import os
import docx

# pede ao usuário o caminho da pasta a ser verificada
dir_path = input("Insira o caminho da pasta a ser verificada: ")

# pede ao usuário o texto a ser buscado e substituído
search_text = input("Insira o texto a ser buscado e substituído: ")
replace_text = input("Insira o novo texto: ")

# itera sobre todos os arquivos .docx na pasta especificada pelo usuário
for filename in os.listdir(dir_path):
    if filename.endswith('.docx'):
        doc_path = os.path.join(dir_path, filename)
        doc = docx.Document(doc_path)
        
        # itera sobre todos os parágrafos do arquivo
        for para in doc.paragraphs:
            if search_text in para.text:
                # substitui o texto encontrado pelo texto inserido pelo usuário
                para.text = para.text.replace(search_text, replace_text)
        
        # itera sobre todas as tabelas do arquivo
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if search_text in cell.text:
                        # substitui o texto encontrado pelo texto inserido pelo usuário
                        cell.text = cell.text.replace(search_text, replace_text)
        
        # salva o arquivo com as modificações
        doc.save(doc_path)
        
print("Substituição concluída com sucesso!")